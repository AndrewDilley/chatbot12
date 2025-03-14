from flask import Flask, request, jsonify, render_template
from openai import OpenAI
from docx import Document
from PyPDF2 import PdfReader
import faiss
import numpy as np
from dotenv import load_dotenv
import os
import json
import faiss


app = Flask(__name__)

# Load environment variables
load_dotenv()

from config import DOCUMENTS_PATH, PREPROCESSED_PATH, SHAREPOINT_LINKS


# Load the FAISS index
index_file_path = os.path.join(PREPROCESSED_PATH, "faiss_index.bin")
index = faiss.read_index(index_file_path)

# Load text map
with open(os.path.join(PREPROCESSED_PATH, "text_map.json"), "r") as f:
    text_map = json.load(f)

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


FILES = list(SHAREPOINT_LINKS.keys())


def split_text_into_chunks(text, chunk_size=8000):
    """
    Splits text into smaller chunks to fit within the token limit.
    """
    chunks = []
    while len(text) > chunk_size:
        split_index = text[:chunk_size].rfind(" ")  # Find the last space to avoid cutting words
        chunks.append(text[:split_index])
        text = text[split_index:].strip()
    chunks.append(text)
    return chunks


# Function to generate embeddings using OpenAI

def generate_embeddings(text):
    """
    Generates embeddings for the given text.
    Splits the text into chunks if it exceeds the token limit.
    """
    embeddings = []
    text_chunks = split_text_into_chunks(text)
    for chunk in text_chunks:
        response = client.embeddings.create(
            input=chunk,
            model="text-embedding-ada-002"
        )
        embeddings.append(response.data[0].embedding)
    return np.mean(embeddings, axis=0)  # Average embedding for simplicity


def search_relevant_text(query, similarity_threshold=0.2, min_length=5):

   # Handle short queries
    if len(query.split()) < min_length:
        return None, "N/A", None


    query_embedding = np.array([generate_embeddings(query)]).astype('float32')
    distances, indices = index.search(query_embedding, k=1)

    print(f"DEBUG: Query: '{query}' | Distances: {distances} | Indices: {indices}")

    # If the best match is below the similarity threshold, return "N/A"
    if distances[0][0] < similarity_threshold:
       return None, "N/A", None

    matched_text, file_name = text_map[indices[0][0]]
    print(f"DEBUG: Matched file at index {indices[0][0]} is {file_name}")

    print(f"DEBUG: Extracted text preview: {matched_text[:200]}...")

    # Extract "Content enquiries" from the document's metadata
    content_enquiries = extract_content_enquiries(file_name)

    return matched_text, file_name, content_enquiries

def extract_content_enquiries(file_name):
    """
    Extract 'Content enquiries' metadata from the document, including adjacent table cells.
    """
    try:
        # Construct the document path
        document_path = os.path.join(DOCUMENTS_PATH, file_name)
        document_path = os.path.normpath(document_path)  # Normalize the path for consistency

        # Check if the file is a .docx
        if not file_name.lower().endswith('.docx'):
            #print(f"Error: Unsupported file format for {file_name}. Only .docx files are supported.")
            return None

        # Verify the file exists
        if not os.path.exists(document_path):
            return None

        # Load the document
        doc = Document(document_path)

        # Search paragraphs first
        for paragraph in doc.paragraphs:
            if "Content enquiries" in paragraph.text:
                return paragraph.text.partition("Content enquiries")[-1].strip()

        # If not found in paragraphs, search tables
        for table in doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    if "Content enquiries" in cell.text:
                        # Ensure there is a next cell in the same row
                        if i + 1 < len(row.cells):
                            adjacent_text = row.cells[i + 1].text.strip()
                            return adjacent_text
                        else:
                            return None

    except Exception as e:
        print(f"Error extracting content enquiries: {e}")
        return None


def generate_response(user_input):
    try:
        # Predefined polite responses
        generic_responses = {"hi", "hello", "hey", "greetings"}
        polite_responses = {"thanks", "thank you", "bye", "goodbye"}

        if user_input.lower() in generic_responses:
            return "Hello! How can I assist you today?"
        elif user_input.lower().strip("!.") in polite_responses:
            return "You're welcome! Let me know if you need anything else!"

        # Get relevant text and file reference
        relevant_text, file_name, content_enquiries = search_relevant_text(user_input)

        if file_name == "N/A":
            prompt = f"You are a helpful assistant. Please answer this question directly:\n\nQuestion: {user_input}"
        else:
            prompt = f"Use the following document text to answer the question:\n\n{relevant_text}\n\nQuestion: {user_input}"

        completion = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ]
        )
        answer = completion.choices[0].message.content


        # Format response into readable sections
        formatted_answer = format_response(answer)

        sharepoint_link = SHAREPOINT_LINKS.get(file_name, "#")

        # Add document reference if applicable
        # Add document reference only if relevant
        if file_name != "N/A":
            formatted_answer = (
                f"{formatted_answer}<br>"
                f"<span class='reference-label'>Reference:</span> "
                f"<a href='{sharepoint_link}' target='_blank' class='reference-text'>{file_name}</a>"
            )
            if content_enquiries:
                formatted_answer += (
                    f"<br><span class='content-enquiries-label'>Content enquiries:</span> "
                    f"<span class='content-enquiries-text'>{content_enquiries}</span>"
                )

        # Reset variables to avoid carry-over
        relevant_text = None
        file_name = None
        content_enquiries = None

        return formatted_answer

    except Exception as e:
        return f"Error: {str(e)}"


import re

def format_response(raw_text):
    """
    Format the raw text into clean HTML:
    1. Convert headings like "### Text" to <h2>.
    2. Convert bullet points (- ...) into <ul><li>.
    3. Convert numbered points (1. **Text**) into <ol><li>.
    4. Remove all instances of ** from the text.
    """
    formatted_text = ""
    lines = raw_text.split("\n")  # Split text by lines

    for line in lines:
        line = line.strip()  # Remove leading/trailing whitespace

        # Skip empty lines
        if not line:
            continue

        # Remove all instances of **
        line = re.sub(r"\*\*", "", line)  # Remove all ** from the text

        # Convert headings (###)
        if line.startswith("#### "):
            formatted_text += f"<h4>{line[5:].strip()}</h4>\n"

        # Convert "### Text" to <h3>
        elif line.startswith("### "):
            formatted_text += f"<h3>{line[4:].strip()}</h3>\n"

        # Convert bullet points (- ...)
        elif line.startswith("- "):
            formatted_text += f"<ul><li>{line[2:].strip()}</li></ul>\n"


        # Normal paragraphs
        else:
            formatted_text += f"<p>{line}</p>\n"

    return formatted_text

# Route for the chatbot
@app.route('/chat', methods=['POST'])
def chat():
    data = request.get_json()
    user_input = data.get('message', '')
    response = generate_response(user_input)
    return jsonify({"response": response})

# Default route
@app.route('/')
def home():
    return render_template('index.html')

# if __name__ == '__main__':
#     if os.getenv("DOCKER_ENV") == "true":
#         app.run(host='0.0.0.0', port=5000, debug=False)
#     else:
#         app.run(host='127.0.0.1', port=5000, debug=True)

# change made on Jan 17 to get the app accessible to other devices on the network

if __name__ == '__main__':
    if os.getenv("DOCKER_ENV") == "true":
        app.run(host='0.0.0.0', port=5000, debug=False)
    else:
        app.run(host='0.0.0.0', port=5000, debug=True)
